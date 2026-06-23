#!/usr/bin/env python3
"""
autotypesetter — turn a messy script manuscript into an InDesign-ready (or
fully designed) Word document with named paragraph styles.

Reads a manuscript in any of these formats:

    .fdx    Final Draft (native XML; element types are declared)
    .rtf    Final Draft-exported RTF
    .docx   Word, including author files that carry no real paragraph styles
    .html   a local HTML file or an http(s) URL (e.g. a Project Gutenberg page)

and writes a .docx in which every paragraph carries a named paragraph style
(Character Name, Dialogue, Stage Directions, Song Lyrics, Scene Number, and so
on). Placed in InDesign with Show Import Options -> Use InDesign Style
Definition, the styles map automatically.

For .fdx and .rtf the element types are read directly. For .docx and .html the
structure is usually absent (everything sits in one default style), so the
parser infers it: a per-document convention pass (centered cues, parenthesized
or italic stage directions, colon / em-dash / period inline cues, indent
levels, musical lyric indent, front matter) followed by tiered, evidence-based
classification with a low-confidence review list.

Output style names are configurable with --config or --map, so the same tool
fits any house style rather than a single fixed preset.

Usage:
    python3 autotypesetter.py script.fdx
    python3 autotypesetter.py manuscript.docx -o out.docx
    python3 autotypesetter.py play.html --start-at-body
    python3 autotypesetter.py "https://www.gutenberg.org/ebooks/12345.html"
    python3 autotypesetter.py musical.fdx --map "Lyrics=***Song Lyrics"
    python3 autotypesetter.py manuscript.docx --report          # diagnostics only
    python3 autotypesetter.py manuscript.docx --template house.docx

Requires: python-docx   (HTML input uses only the Python standard library.)
"""

import argparse
import json
import os
import re
import sys
import xml.etree.ElementTree as ET
from collections import Counter

# ---------------------------------------------------------------------------
# Default mapping: Final Draft element type  ->  InDesign paragraph style name
# Anything not listed passes through under "FD <Type>" so nothing is lost
# and unmapped elements are easy to spot in the Paragraph Styles panel.
# ---------------------------------------------------------------------------
DEFAULT_STYLE_MAP = {
    "Action":        "***Stage Directions",
    "Character":     "***Character Name",
    "Dialogue":      "***Dialogue",
    "Transition":    "***Stage Directions",
    "Lyrics":        "***Song Lyrics",
    "Song Title":    "***Song Title",
    "Scene Heading": "***Scene Number",
    "Simultaneous":  "***Simultaneous",
    "General":       "FD General (front matter)",
    # "Parenthetical" is intentionally unmapped by default -> "FD Parenthetical"
    # (decide per show: inline italics vs. its own style). Override with
    # e.g. --map "Parenthetical=***Dialogue"
}

# Default mapping for --template (designed Word output): internal element type
# -> the donor template's paragraph style name. Overridable via a config file.
DEFAULT_TEMPLATE_MAP = {
    "Character":     "CHARACTER NAME",
    "Dialogue":      "Dialogue",
    "Lyrics":        "LYRICS",
    "Action":        "Stage Direction",
    "Song Title":    "Song Titles",
    "Scene Heading": "SCENE",
    "Simultaneous":  "Dialogue",
}


def load_config(path):
    """Load element->style mappings from a JSON config file. The file may
    contain an "indesign" object, a "template" object, or both; either is
    merged over the built-in defaults, so a partial file only overrides the
    keys it names. Returns (indesign_map, template_map)."""
    try:
        with open(path, encoding="utf-8") as fh:
            data = json.load(fh)
    except FileNotFoundError:
        sys.exit(f"error: config file not found: {path}")
    except json.JSONDecodeError as e:
        sys.exit(f"error: config file is not valid JSON ({path}): {e}")
    if not isinstance(data, dict):
        sys.exit(f"error: config root must be a JSON object: {path}")

    indesign = dict(DEFAULT_STYLE_MAP)
    template = dict(DEFAULT_TEMPLATE_MAP)
    # Accept either nested form {"indesign": {...}, "template": {...}} or a
    # flat object, which is treated as InDesign overrides for convenience.
    if "indesign" in data or "template" in data:
        for section, target in (("indesign", indesign), ("template", template)):
            block = data.get(section, {})
            if not isinstance(block, dict):
                sys.exit(f'error: "{section}" in config must be an object: {path}')
            for k, v in block.items():
                if not isinstance(v, str):
                    sys.exit(f'error: mapping value for "{k}" must be a string')
                target[k] = v
    else:
        for k, v in data.items():
            if not isinstance(v, str):
                sys.exit(f'error: mapping value for "{k}" must be a string')
            indesign[k] = v
    return indesign, template



UPPERCASE_TYPES = {"Character"}  # FD displays these in caps; the file stores typed text


# ---------------------------------------------------------------------------
# Shared paragraph model: {"type": str, "runs": [{"t": str, "i","b","u": bool}]}
# ---------------------------------------------------------------------------

class _SegRun:
    """A run-like view over a pre-split segment dict {"t","i","b","u"}."""
    __slots__ = ("text", "italic", "bold", "underline")

    def __init__(self, part):
        self.text = part["t"]
        self.italic = part["i"]
        self.bold = part["b"]
        self.underline = part["u"]


class _SegPara:
    """A paragraph-like wrapper around one visual line (segment) of a docx
    paragraph that contained soft line breaks. Inherits the parent paragraph's
    alignment, indentation, and style so all classification signals still work,
    but exposes only this segment's runs and text."""
    def __init__(self, seg_runs, parent, force_type=None):
        self._runs = [_SegRun(p) for p in seg_runs if p["t"] != ""]
        self.alignment = parent.alignment
        self.paragraph_format = parent.paragraph_format
        self.style = parent.style
        self._p = parent._p
        self.force_type = force_type

    @property
    def runs(self):
        return self._runs

    @property
    def text(self):
        return "".join(r.text for r in self._runs)


def _merge_run(runs, text, i=False, b=False, u=False):
    if not text:
        return
    if runs and runs[-1]["i"] == i and runs[-1]["b"] == b and runs[-1]["u"] == u:
        runs[-1]["t"] += text
    else:
        runs.append({"t": text, "i": i, "b": b, "u": u})


def _finish(runs):
    """Strip leading and trailing whitespace (incl. author tab-indents, which
    are junk in InDesign where paragraph styles own indentation); return None
    if nothing remains."""
    while runs and runs[-1]["t"].rstrip() == "":
        runs.pop()
    while runs and runs[0]["t"].lstrip() == "":
        runs.pop(0)
    if not runs:
        return None
    runs[0]["t"] = runs[0]["t"].lstrip()
    runs[-1]["t"] = runs[-1]["t"].rstrip()
    return runs


# ---------------------------------------------------------------------------
# FDX parser (native Final Draft XML — the preferred input)
# ---------------------------------------------------------------------------

def parse_fdx(path):
    tree = ET.parse(path)
    content = tree.getroot().find("Content")
    if content is None:
        sys.exit("error: no <Content> element — is this a Final Draft .fdx file?")

    paras, dual_blocks = [], []

    def emit(p):
        ptype = p.get("Type") or "General"
        runs = []
        for tx in p.findall("Text"):
            style = tx.get("Style") or ""
            _merge_run(
                runs,
                tx.text or "",
                i="Italic" in style,
                b="Bold" in style,
                u="Underline" in style,
            )
        runs = _finish(runs)
        if runs:
            paras.append({"type": ptype, "runs": runs})

    for p in content.findall("Paragraph"):
        dd = p.find("DualDialogue")
        if dd is not None:
            first = dd.find("Paragraph/Text")
            dual_blocks.append(
                (len(paras), (first.text or "").strip() if first is not None else "?")
            )
            for sub in dd.findall("Paragraph"):
                emit(sub)
            continue
        emit(p)

    return paras, dual_blocks


# ---------------------------------------------------------------------------
# FD-exported RTF parser (fallback when no .fdx is available)
#
# Final Draft's RTF export contains no RTF stylesheet; element identity lives
# in {\*\FDElementName:Type,start,end} metadata groups whose offsets index
# into the document's plain-text stream. The text is extracted (tracking
# italic/bold/underline state per character) and sliced by those ranges.
# ---------------------------------------------------------------------------

_DEST_WORDS = {
    "fonttbl", "colortbl", "expandedcolortbl", "stylesheet",
    "info", "pict", "header", "footer",
}


def parse_fd_rtf(path):
    raw = open(path, encoding="latin-1").read()

    elements = [
        (m.group(1), int(m.group(2)), int(m.group(3)))
        for m in re.finditer(r"\{\\\*\\FDElementName:([^,}]+),(\d+),(\d+)\}", raw)
    ]
    if not elements:
        sys.exit(
            "error: no FDElementName metadata found — this RTF was not exported "
            "by Final Draft (or was re-saved by another app, stripping the tags). "
            "Ask for the .fdx instead."
        )

    i, n = 0, len(raw)
    chars = []                      # (char, italic, bold, underline)
    stack, italic, bold, under = [], False, False, False

    def emit(c):
        chars.append((c, italic, bold, under))

    def skip_group(j):
        depth = 1
        while j < n and depth:
            if raw[j] == "\\" and j + 1 < n:
                j += 2
                continue
            if raw[j] == "{":
                depth += 1
            elif raw[j] == "}":
                depth -= 1
            j += 1
        return j

    while i < n:
        c = raw[i]
        if c == "{":
            stack.append((italic, bold, under))
            i += 1
            if raw.startswith("\\*", i):
                i = skip_group(i + 2)
                stack.pop()
                continue
            m = re.match(r"\\([a-z]+)", raw[i:])
            if m and m.group(1) in _DEST_WORDS:
                i = skip_group(i)
                stack.pop()
            continue
        if c == "}":
            if stack:
                italic, bold, under = stack.pop()
            i += 1
            continue
        if c == "\\":
            nxt = raw[i + 1] if i + 1 < n else ""
            if nxt in "\n\r":                       # \<newline> = paragraph break
                emit("\n")
                i += 2
                if nxt == "\r" and i < n and raw[i] == "\n":
                    i += 1
                continue
            if nxt in "\\{}":
                emit(nxt)
                i += 2
                continue
            if nxt == "'":                          # \'xx hex escape (cp1252)
                try:
                    emit(bytes([int(raw[i + 2:i + 4], 16)]).decode("cp1252"))
                except ValueError:
                    emit("?")
                i += 4
                continue
            if nxt == "~":
                emit("\u00a0"); i += 2; continue
            if nxt == "_":
                emit("\u2011"); i += 2; continue
            m = re.match(r"\\([a-z]+)(-?\d+)? ?", raw[i:])
            if m:
                word, param = m.group(1), m.group(2)
                if word in ("par", "line"):
                    emit("\n")
                elif word == "tab":
                    emit("\t")
                elif word == "i":
                    italic = param != "0"
                elif word == "b":
                    bold = param != "0"
                elif word in ("ul", "ulnone"):
                    under = word == "ul" and param != "0"
                elif word == "u" and param is not None:
                    cp = int(param)
                    emit(chr(cp + 65536 if cp < 0 else cp))
                i += m.end()
                continue
            i += 2
            continue
        if c in "\r\n":
            i += 1
            continue
        emit(c)
        i += 1

    expected = max(e[2] for e in elements)
    if len(chars) != expected:
        print(
            f"warning: extracted {len(chars)} chars but FD metadata expects "
            f"{expected}; element boundaries may drift. Prefer the .fdx.",
            file=sys.stderr,
        )

    paras = []
    for name, s, e in elements:
        line = []

        def flush():
            nonlocal line
            runs = []
            for ch, it, b, u in line:
                _merge_run(runs, ch, i=it, b=b, u=u)
            runs = _finish(runs)
            if runs:
                paras.append({"type": name, "runs": runs})
            line = []

        for tup in chars[s:e]:
            if tup[0] == "\n":
                flush()
            else:
                line.append(tup)
        flush()

    return paras, []   # RTF export already flattens dual dialogue


# ---------------------------------------------------------------------------
# Word (.docx) parser — for scripts authors typed in Word
#
# Three tiers of evidence, best first:
#   1. Real paragraph style names, if the author used any (rare but reliable)
#   2. Formatting/text patterns: scene headings, ALL-CAPS cues, italics,
#      parenthesized directions, "NAME: dialogue" inline cues
#   3. Context fallback (a paragraph after a character cue is dialogue) —
#      these are flagged as low-confidence in the report
# ---------------------------------------------------------------------------

# Author style names recognized, normalized -> FD element type
_WORD_STYLE_ALIASES = {
    "character": "Character", "character name": "Character",
    "dialogue": "Dialogue", "dialog": "Dialogue",
    "action": "Action", "stage direction": "Action", "stage directions": "Action",
    "parenthetical": "Parenthetical",
    "lyrics": "Lyrics", "lyric": "Lyrics", "song lyrics": "Lyrics",
    "scene heading": "Scene Heading", "scene": "Scene Heading",
    "transition": "Transition",
}

_SCENE_RE = re.compile(
    r"^(ACT\s+[\dIVXLC]+|SCENE\s+[\dIVXLC]+|PROLOGUE|EPILOGUE|"
    r"Act\s+(One|Two|Three|Four|Five|[\dIVXLC]+)|Scene\s+\d+)\b[.:]?\s*$"
)
_INLINE_CUE_RE = re.compile(r"^([A-Z][A-Z0-9 .'\-&/]{0,30}?)\s*:\s+(\S.*)$")
# Capitalized (not necessarily ALL-CAPS) short label + colon + text, e.g.
# "Anna: ..." or "The Witness: ...". 1-3 words, each starting uppercase.
_INLINE_LABEL_RE = re.compile(
    r"^([A-Z][\w'\-]*(?:\s+[A-Z][\w'\-]*){0,2})\s*:\s+(\S.*)$")
# Inline cue using an EM-DASH separator instead of a colon — the dominant
# convention in older published editions (Gutenberg etc.):
#   "Anna—Why is the gate still open?"
#   "Mr. Vance—[turning away] Not tonight."
# 1-3 Capitalized words (allowing internal periods for "Mrs."/"Dr."), then a
# literal em-dash (U+2014, NOT a hyphen or "--"), then the speech.
_INLINE_DASH_RE = re.compile(
    r"^([A-Z][\w'’.\-]*(?:\s+[A-Z][\w'’.\-]*){0,2})\s*\u2014\s*(\S.*)$")
# Inline cue using a PERIOD after an ALL-CAPS name — Shaw and many Gutenberg
# play editions:  "ANNA. Not a word more!"   "THE COMMANDER. Hold the line."
# 1-4 ALL-CAPS words then a period then the speech. (Caps-ness is enforced in
# the detector, not the regex, so the captured label can be reused.)
_INLINE_PERIOD_RE = re.compile(
    r"^([A-Z][A-Z0-9'’.\-&]*(?:\s+[A-Z][A-Z0-9'’.\-&]*){0,3})\.\s+(\S.*)$")
# Same naming family, but the name is followed immediately by a bracketed
# stage direction instead of a period:  "ANNA [coldly] You are mistaken."
_INLINE_CAPS_DIR_RE = re.compile(
    r"^([A-Z][A-Z0-9'’.\-&]*(?:\s+[A-Z][A-Z0-9'’.\-&]*){0,3})\s+([\[(].*)$")
# A leading bracketed/parenthesized stage direction at the head of a speech,
# e.g. "[rising] Listen to me." -> peel "[rising]" off as Action.
_LEAD_DIR_RE = re.compile(r"^([\[(][^\])]*[\])])\s*(.*)$")
# Bare bracketed page markers from scanned/ebook editions: "[8]", "[iv]". These
# are not script content and are dropped entirely.
_PAGE_NUM_RE = re.compile(r"^\[(?:\d+[a-z]?|[ivxlcdm]+)\]$", re.I)
# A scene heading carrying its set description on the same line, em-dash joined:
# "Scene I.—A small room at dusk." -> Scene Number "Scene I" + Stage Direction.
# (A bare "SCENE 1" with no inline text is handled by _SCENE_RE instead.)
_SCENE_DESC_RE = re.compile(
    r"^((?:ACT|SCENE)\s+[\dIVXLC]+|(?:Act|Scene)\s+[\dIVXLC]+)\.?\s*[—–]\s*(\S.*)$")
_CUE_TAIL_RE = re.compile(r"\s*(\(.*\)|\[.*\])\s*$")   # (CONT'D), (offstage)...

# Short ALL-CAPS lines that are stagecraft, not character cues
_STAGE_TERMS_RE = re.compile(
    r"^(BLACKOUT|BLACK OUT|CURTAIN|INTERMISSION|"
    r"LIGHTS\s+(UP|DOWN|OUT|FADE|SHIFT|RISE).*|"
    r"(END|TOP)\s+OF\s+(ACT|SCENE|PLAY|SHOW).*|"
    r"FADE\s+(IN|OUT|TO\s+BLACK).*|THE\s+END)[.!]?\s*$"
)


def _is_caps_cue(text):
    """Short ALL-CAPS line with no sentence-ending punctuation = character cue.
    Tolerates compound cues like 'ANNA and MARCO' / 'ANNA & MARCO'."""
    core = _CUE_TAIL_RE.sub("", text).strip().rstrip(":").strip()
    if not core or len(core) > 40:
        return False
    # ignore lowercase connector words when judging caps-ness
    core_check = re.sub(r"\b(and|und|y|et|&|or|the)\b", " ", core, flags=re.I)
    letters = [c for c in core_check if c.isalpha()]
    if not letters:
        return False
    if sum(c.isupper() for c in letters) / len(letters) < 0.95:
        return False
    return not core.endswith((".", "!", "?", ","))


def _italic_ratio(para):
    tot = ital = 0
    for r in para.runs:
        n = len(r.text)
        tot += n
        if r.italic or (r.italic is None and getattr(r.style, "font", None)
                        and r.style.font.italic):
            ital += n
    return (ital / tot) if tot else 0.0


def _bold_ratio(para):
    tot = bold = 0
    for r in para.runs:
        n = len(r.text)
        tot += n
        if r.bold:
            bold += n
    return (bold / tot) if tot else 0.0


def _is_centered(para):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return True
    try:
        return para.style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    except AttributeError:
        return False


def _left_indent(para):
    """Left indent in inches (0.0 if none), rounded to handle sub-point noise."""
    li = para.paragraph_format.left_indent
    if li is None:
        return 0.0
    try:
        return round(li.inches, 2)
    except (AttributeError, ValueError):
        return 0.0


def _round_indent(v, step=0.1):
    """Bucket an indent to the nearest `step` so 0.79/0.80/0.81 group together."""
    return round(round(v / step) * step, 2)


def _eff_indent(para):
    """Effective visual left indent in inches, used for LYRIC detection only.
    Authors indent lyrics three different ways and they should all count:
    a real paragraph left-indent, a positive first-line indent, or one or more
    leading tab characters (~0.5" each). `_left_indent` alone misses the latter
    two — e.g. a musical whose lyrics are first-line-indented or tab-indented
    reads as indent 0 and never trips the musical heuristic. This does NOT
    replace `_left_indent` (which the dialogue/direction indent convention still
    uses); it is a lyric-specific signal."""
    li = _left_indent(para)
    fli = para.paragraph_format.first_line_indent
    try:
        fli = max(fli.inches, 0.0) if fli else 0.0
    except (AttributeError, ValueError):
        fli = 0.0
    raw = "".join(r.text for r in para.runs)
    tabs = len(raw) - len(raw.lstrip("\t"))
    return li + fli + tabs * 0.5


def _is_allcaps_line(text):
    """True if the alphabetic content of the line is essentially all caps."""
    letters = [c for c in text if c.isalpha()]
    if not letters:
        return False
    return sum(c.isupper() for c in letters) / len(letters) >= 0.9


# A song-number/title line: "12.  THE LETTER" or "2.\tOPENING NUMBER".
_SONG_TITLE_RE = re.compile(r"^\s*\d+\.\s*\t?\s*(\S.*)$")


def _split_cue_names(text):
    """Split a cue line into individual character names, handling shared cues
    like 'ANNA AND MARCO', 'ANNA, MARCO, AND LEE', 'MARCO & LEE'.
    Returns a list of upper-cased name tokens (may be multi-word like 'OLD JOE')."""
    t = re.sub(r"[,&]", " ", text)
    t = re.sub(r"\bAND\b", " ", t)
    # collapse whitespace; treat remaining runs of words as names. For the
    # common single-word-name case this yields clean tokens; multiword names
    # (OLD JOE) are kept whole only when not joined by a separator.
    parts = [p.strip() for p in t.split() if p.strip()]
    return parts


def _docx_runs(para, force_text=None):
    runs = []
    if force_text is not None:
        _merge_run(runs, force_text)
        return _finish(runs)
    for r in para.runs:
        _merge_run(runs, r.text, i=bool(r.italic), b=bool(r.bold),
                   u=bool(r.underline))
    return _finish(runs)


def _emit_cue(paras, para, label, remainder):
    """Emit a split inline cue: a Character paragraph for `label`, an Action
    paragraph for any leading bracketed/parenthesized direction peeled off
    `remainder`, then a Dialogue paragraph for the rest. A mid-line direction is
    left inline (splitting it reliably would break the sentence). Returns the
    resulting prev_type."""
    runs = _docx_runs(para, force_text=label)
    if runs:
        paras.append({"type": "Character", "runs": runs})
    dm = _LEAD_DIR_RE.match(remainder)
    if dm:
        d_runs = _docx_runs(para, force_text=dm.group(1))
        if d_runs:
            paras.append({"type": "Action", "runs": d_runs})
        remainder = dm.group(2).strip()
    if remainder:
        d_runs = _docx_runs(para, force_text=remainder)
        if d_runs:
            paras.append({"type": "Dialogue", "runs": d_runs})
        return "Dialogue"
    return "Action"


def _split_runs_on_breaks(para):
    """Yield one run-list per manual-line-break-separated segment."""
    segment = []
    for r in para.runs:
        i, b, u = bool(r.italic), bool(r.bold), bool(r.underline)
        parts = re.split(r"[\n\x0b]", r.text)
        for k, part in enumerate(parts):
            if k > 0:                       # a break occurred before this part
                yield segment
                segment = []
            _merge_run(segment, part, i=i, b=b, u=u)
    yield segment


def parse_word_docx(path, breaks="split", keep_blanks=False):
    try:
        from docx import Document
    except ImportError:
        sys.exit("error: python-docx not installed — run:  pip install python-docx")

    doc = Document(path)

    # --- Convention-detection pass -------------------------------------
    # If the author centers character cues (common), alignment splits short
    # ALL-CAPS lines into cues (centered, not bold) vs. scene/section title
    # cards (bold, or left-aligned).
    caps_centered = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t and _is_caps_cue(t):
            caps_centered.append(_is_centered(para))
    cues_centered = (
        len(caps_centered) >= 20
        and sum(caps_centered) / len(caps_centered) >= 0.5
    )

    # Does the author wrap ALL stage directions in parentheses/brackets? If so,
    # a non-wrapped line mid-speech is reliably dialogue, and speech context can
    # carry across interjected directions. Detected by checking whether
    # italic-or-isolated direction-like lines are almost all wrapped.
    wrapped_dirs = unwrapped_dirs = 0
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t or _is_caps_cue(t):
            continue
        is_wrapped = (t.startswith("(") and t.endswith(")")) or \
                     (t.startswith("[") and t.endswith("]"))
        if is_wrapped:
            wrapped_dirs += 1
        elif _italic_ratio(para) > 0.6:
            unwrapped_dirs += 1
    parens_convention = (
        wrapped_dirs >= 15
        and wrapped_dirs / max(1, wrapped_dirs + unwrapped_dirs) >= 0.9
    )

    # Does the author write directions as italic text (rather than parens)?
    # Italic directions must be both numerous AND the dominant form (most
    # non-cue, non-dialogue-ish lines are italic), so plays that merely use
    # italics for emphasis do not trigger it. "Looks like a direction" is
    # approximated as: not a cue, not wrapped, sentence-like prose.
    italic_dirs = plain_prose = 0
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t or _is_caps_cue(t):
            continue
        wrapped = (t.startswith("(") and t.endswith(")")) or \
                  (t.startswith("[") and t.endswith("]"))
        if wrapped:
            continue
        if _italic_ratio(para) > 0.6:
            italic_dirs += 1
        else:
            plain_prose += 1
    # Italics dominate the non-cue prose only if italic lines clearly outnumber
    # plain ones. In a "Name: dialogue" play most plain lines are cues+speech,
    # so the downstream rule is guarded with a ratio test rather than a bare
    # count.
    italic_directions = (
        italic_dirs >= 15
        and italic_dirs >= plain_prose * 0.6
    )

    # Does the author label speakers inline as "Name: dialogue" on one line
    # (e.g. "Anna: I can manage this"), with the name Capitalized but not
    # necessarily ALL-CAPS, and inconsistent tabbing? Counted by matching a
    # short Capitalized label + colon + text where the same labels recur (real
    # speakers repeat; prose colons do not).
    from collections import Counter as _C
    label_counts = _C()
    inline_candidate_lines = 0
    for para in doc.paragraphs:
        t = para.text.strip()
        m = _INLINE_LABEL_RE.match(t)
        if m and _italic_ratio(para) <= 0.6:
            label_counts[m.group(1).strip()] += 1
            inline_candidate_lines += 1
    # require: enough candidate lines, and the labels are a small recurring set
    recurring = {lbl for lbl, n in label_counts.items() if n >= 2}
    inline_cue_convention = (
        inline_candidate_lines >= 8
        and len(recurring) >= 1
        and sum(label_counts[l] for l in recurring) >= inline_candidate_lines * 0.5
    )
    inline_cue_labels = recurring

    # Does the author label speakers inline with an EM-DASH instead of a colon
    # ("Anna—Why...")? Same recurrence logic as the colon case, plus guards
    # against the two shapes that mimic a name-dash but are not cues:
    #   * scene headings:  "Scene I.—A small room at dusk..."
    #   * front-matter notices/abbreviations ending in a period: "CAUTION.—..."
    # A real speaker label does not end in a period and is not Scene/Act/etc.
    dash_label_counts = _C()
    inline_dash_candidates = 0
    for para in doc.paragraphs:
        t = para.text.strip()
        m = _INLINE_DASH_RE.match(t)
        if not m or _italic_ratio(para) > 0.6:
            continue
        label = m.group(1).strip()
        if label.endswith(".") or len(label) > 28:
            continue
        if re.match(r"(?i)^(scene|act|prologue|epilogue|curtain)\b", label):
            continue
        dash_label_counts[label] += 1
        inline_dash_candidates += 1
    dash_recurring = {lbl for lbl, n in dash_label_counts.items() if n >= 2}
    inline_dash_convention = (
        inline_dash_candidates >= 8
        and len(dash_recurring) >= 1
        and sum(dash_label_counts[l] for l in dash_recurring)
            >= inline_dash_candidates * 0.5
    )
    inline_dash_labels = dash_recurring

    # Does the author put an ALL-CAPS speaker name followed by a period on the
    # same line as the speech ("ANNA. Not a word more!")? Same recurrence logic; the
    # name must be ALL-CAPS (so prose "Mr. Smith. " and sentences do not match)
    # and not a scene/act word. Speeches whose name is followed by a bracketed
    # direction instead of a period ("ANNA [coldly] Enough!") count too, using
    # the same recurring-speaker set.
    period_label_counts = _C()
    inline_period_candidates = 0
    for para in doc.paragraphs:
        t = para.text.strip()
        if _italic_ratio(para) > 0.6:
            continue
        m = _INLINE_PERIOD_RE.match(t)
        if not m:
            continue
        label = m.group(1).strip()
        if not _is_allcaps_line(label) or len(label) > 30:
            continue
        if re.match(r"(?i)^(scene|act|prologue|epilogue|curtain)\b", label):
            continue
        period_label_counts[label] += 1
        inline_period_candidates += 1
    period_recurring = {lbl for lbl, n in period_label_counts.items() if n >= 2}
    inline_period_convention = (
        inline_period_candidates >= 8
        and len(period_recurring) >= 1
        and sum(period_label_counts[l] for l in period_recurring)
            >= inline_period_candidates * 0.5
    )
    inline_period_labels = period_recurring

    # Does the author distinguish dialogue from stage directions by LEFT INDENT?
    # Many manuscripts indent dialogue deeper than directions (or vice versa)
    # consistently, even when italics are erratic. Among non-cue (non-flush,
    # non-ALL-CAPS) lines, if two clear indent levels dominate, the deeper level
    # is dialogue and the shallower is stage directions. Often the most reliable
    # signal when present.
    indent_hist = _C()
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t or _is_caps_cue(t):
            continue
        li = _round_indent(_left_indent(para))
        if li > 0.05:                     # ignore flush-left (cues / prose)
            indent_hist[li] += 1
    indent_convention = False
    dialogue_indent = direction_indent = None
    if len(indent_hist) >= 2:
        top = indent_hist.most_common(2)
        (lvl_a, n_a), (lvl_b, n_b) = top[0], top[1]
        # the two dominant indent levels must each be substantial and well
        # separated, and together cover most indented lines
        total_indented = sum(indent_hist.values())
        if (n_a >= 10 and n_b >= 8
                and abs(lvl_a - lvl_b) >= 0.2
                and (n_a + n_b) >= total_indented * 0.8):
            indent_convention = True
            dialogue_indent = max(lvl_a, lvl_b)     # deeper = dialogue
            direction_indent = min(lvl_a, lvl_b)    # shallower = direction

    # ----- Musical detection -------------------------------------------------
    # In a musical, lyrics are typically set as ALL-CAPS lines at a consistent
    # indent, distinct from spoken dialogue (mixed-case, usually flush-left).
    # Song titles appear as a numbered, bold line ("12.  THE LETTER"). Detected:
    #   * the lyric indent: the indent level where ALL-CAPS lines pile up
    #   * the cast: recurring bare ALL-CAPS names (used to tell a singing CUE
    #     at the lyric indent from a LYRIC line at the same indent)
    caps_indent_hist = _C()
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t or not _is_allcaps_line(t):
            continue
        li = _round_indent(_eff_indent(para))
        if li > 0.05:
            caps_indent_hist[li] += 1
    lyric_indent = None
    is_musical = False
    if caps_indent_hist:
        cand, n = caps_indent_hist.most_common(1)[0]
        if n >= 40:                       # a real musical has many lyric lines
            lyric_indent = cand
            is_musical = True
    # Build a cast set from short recurring flush-left ALL-CAPS names, to
    # distinguish a singing cue from a lyric at lyric_indent.
    cast_names = set()
    if is_musical:
        name_counts = _C()
        for para in doc.paragraphs:
            t = para.text.strip()
            # only flush-left short ALL-CAPS lines are unambiguous spoken cues;
            # lyric-indent lines are also caps and would pollute the cast set
            li = _round_indent(_eff_indent(para))
            if (t and li <= 0.05 and _is_allcaps_line(t) and len(t) <= 30
                    and not _SONG_TITLE_RE.match(t)):
                for nm in _split_cue_names(t):
                    name_counts[nm] += 1
        cast_names = {nm for nm, c in name_counts.items()
                      if c >= 2 and len(nm) >= 2}

    paras, guessed = [], []
    prev_type = None
    in_speech = False     # currently inside a character's speech block?
    song_mode = False     # currently inside a musical number?

    # Some manuscripts bundle a cue and its dialogue into ONE paragraph with a
    # soft line break, e.g. "ANNA<break>It looks like". The combined text
    # classifies as neither a clean cue nor clean dialogue, so such paragraphs
    # are expanded up front into single-line "logical paragraphs" — but ONLY
    # when the first visual line is an ALL-CAPS cue, leaving ordinary multi-line
    # dialogue/directions to the normal break-splitting path below.
    def _is_column_para(raw):
        # A simultaneous-dialogue column row uses runs of consecutive tabs as
        # gutters between 2+ short segments (e.g. "ANNA\t\t\tMARCO\t\t\tLEE").
        # Guard against stray single tabs dropped mid-sentence (long prose on
        # both sides) by requiring (a) at least one gutter of 2+ tabs and
        # (b) segments that are short-ish, not full sentences.
        if "\t\t" not in raw:
            return False
        # a leading run of tabs is indentation, not a column gutter — strip it
        # before analysing, and reject parenthesized directions outright
        body = raw.lstrip("\t")
        if body[:1] in ("(", "["):
            return False
        if "\t\t" not in body:
            return False
        segs = [s.strip() for s in re.split(r"\t+", body) if s.strip()]
        if len(segs) < 2:
            return False
        if max(len(s) for s in segs) > 40:
            return False
        return True

    def _logical_paragraphs(doc):
        for para in doc.paragraphs:
            raw = "".join(r.text for r in para.runs)
            if _is_column_para(raw):
                # pass the whole column block through untouched, tagged so the
                # writer gives it the ***Simultaneous style (tabs/breaks kept)
                yield _SegPara([{"t": raw, "i": False, "b": False, "u": False}],
                               para, force_type="Simultaneous")
                continue
            if re.search(r"[\n\x0b]", raw):
                segs = list(_split_runs_on_breaks(para))
                seg_texts = ["".join(part["t"] for part in s) if s else ""
                             for s in segs]
                first = seg_texts[0].strip() if seg_texts else ""
                nonempty = [t for t in seg_texts if t.strip()]
                if len(nonempty) >= 2 and _is_caps_cue(first):
                    for s in segs:
                        if s and any(part["t"].strip() for part in s):
                            yield _SegPara(s, para)
                    continue
            yield para

    # --- Front-matter detection -----------------------------------------
    # A title page, cast list (e.g. "ANNA: the eldest sister") and a few
    # production notes commonly precede the first scene. Run through the body
    # classifier they misclassify: cast entries split into Character+Dialogue,
    # notes become Dialogue, a bold title becomes a Scene Number. Instead, the
    # preamble is detected and tagged "General" (the front-matter style). The
    # preamble is everything before the first scene heading, but only when that
    # heading appears early — a title-page-sized run, not a long introductory
    # essay — so real body text is never swallowed (a Gutenberg edition with a
    # multi-page intro before ACT I is left alone).
    logical = list(_logical_paragraphs(doc))
    FRONTMATTER_MAX = 40          # max non-empty preamble lines before the body
    body_start_i = None
    seen_nonempty = 0
    for i, p in enumerate(logical):
        t = p.text.strip()
        if not t:
            continue
        if _SCENE_RE.match(t):
            body_start_i = i
            break
        seen_nonempty += 1
        if seen_nonempty > FRONTMATTER_MAX:
            break
    has_frontmatter = body_start_i is not None and 1 <= seen_nonempty <= FRONTMATTER_MAX

    for idx, para in enumerate(logical):
        text = para.text.strip()

        # Front-matter gate: everything before the first (early) scene heading
        # is title/cast/notes, not script body. Emit it as General so it maps
        # to the front-matter style and never mixes into Character/Dialogue/
        # Scene. Checked before blank handling so preamble blanks flow normally.
        if has_frontmatter and text and idx < body_start_i:
            runs = _docx_runs(para)
            if runs:
                paras.append({"type": "General", "runs": runs})
            prev_type = None
            in_speech = False
            continue

        if not text:
            # Under a reliable parens convention, a speaker's block runs until
            # the next cue or scene heading; blank lines within it (including
            # between the cue and its opening direction) do not end it. Without
            # the convention, a blank ends the block unless only the cue has
            # been seen so far.
            if not parens_convention and prev_type != "Character":
                in_speech = False
            # Preserve the author's vertical spacing when requested: emit a
            # blank marker the template writer turns into an empty paragraph.
            if keep_blanks and paras and paras[-1]["type"] != "Blank":
                paras.append({"type": "Blank", "runs": []})
            prev_type = None
            continue

        # Drop bare ebook/scan page markers ("[8]", "[iv]") — not script content.
        if _PAGE_NUM_RE.match(text):
            continue

        style_name = (para.style.name or "").strip().lower().lstrip("*").strip()
        ptype, confident = None, True

        # Scene heading with an inline set description on the same line
        # ("Scene I.—A small room at dusk."): emit the heading as a Scene
        # Number and the description as a Stage Direction. Checked early so the
        # em-dash cue splitter never mistakes the heading for a speaker.
        if style_name not in _WORD_STYLE_ALIASES:
            m_sd = _SCENE_DESC_RE.match(text)
            if m_sd:
                runs = _docx_runs(para, force_text=m_sd.group(1).strip())
                if runs:
                    paras.append({"type": "Scene Heading", "runs": runs})
                runs = _docx_runs(para, force_text=m_sd.group(2).strip())
                if runs:
                    paras.append({"type": "Action", "runs": runs})
                prev_type = "Action"
                in_speech = False
                continue

        # Forced type: a simultaneous-dialogue column block, passed through
        # whole with tabs/breaks intact so it can be rebuilt in InDesign.
        if getattr(para, "force_type", None) == "Simultaneous":
            runs = [{"t": para.text, "i": False, "b": False, "u": False}]
            paras.append({"type": "Simultaneous", "runs": runs})
            prev_type = None
            in_speech = False
            continue

        # Tier 1: author actually used named styles
        if style_name in _WORD_STYLE_ALIASES:
            ptype = _WORD_STYLE_ALIASES[style_name]

        # Tier 1.5: musical-aware classification
        if ptype is None and is_musical:
            li = _round_indent(_eff_indent(para))
            # song-title line: numbered + bold (e.g. "12.  THE LETTER")
            m_song = _SONG_TITLE_RE.match(text)
            if (m_song and _bold_ratio(para) > 0.5
                    and _is_allcaps_line(m_song.group(1))):
                ptype = "Song Title"
                song_mode = True
            elif lyric_indent is not None and abs(li - lyric_indent) <= 0.08 \
                    and _is_allcaps_line(text):
                # at the lyric indent, an ALL-CAPS line is a LYRIC unless it is
                # a bare character name (a singing cue)
                names = _split_cue_names(text)
                # a singing cue is a line whose tokens are ALL cast names
                # (single "ANNA" or shared "ANNA, MARCO, AND LEE"); length
                # is not a reliable guard for full-company cues, so the test
                # is "all tokens are cast names" plus a cap on token count.
                is_cue = (names and len(names) <= 8
                          and all(n in cast_names for n in names))
                if is_cue:
                    ptype = "Character"
                else:
                    ptype = "Lyrics"
                    song_mode = True

        # Tier 2: pattern evidence
        if ptype is None and _SCENE_RE.match(text):
            ptype = "Scene Heading"

        if ptype is None and _STAGE_TERMS_RE.match(text):
            ptype = "Action"

        if ptype is None:
            m = _INLINE_CUE_RE.match(text)
            if m and _is_caps_cue(m.group(1)):
                # "ANNA: I need to be there." -> split into cue + dialogue
                runs = _docx_runs(para, force_text=m.group(1))
                if runs:
                    paras.append({"type": "Character", "runs": runs})
                runs = _docx_runs(para, force_text=m.group(2))
                if runs:
                    paras.append({"type": "Dialogue", "runs": runs})
                prev_type = "Dialogue"
                in_speech = True
                continue

        if ptype is None and inline_cue_convention and _italic_ratio(para) <= 0.6:
            m = _INLINE_LABEL_RE.match(text)
            if m and m.group(1).strip() in inline_cue_labels:
                # "Anna: I can manage this" -> Capitalized inline cue + dialogue.
                # Gated on a recurring known speaker label so prose colons
                # ("She thought: maybe") don't trigger it.
                runs = _docx_runs(para, force_text=m.group(1).strip())
                if runs:
                    paras.append({"type": "Character", "runs": runs})
                runs = _docx_runs(para, force_text=m.group(2))
                if runs:
                    paras.append({"type": "Dialogue", "runs": runs})
                prev_type = "Dialogue"
                in_speech = True
                continue

        if ptype is None and inline_dash_convention and _italic_ratio(para) <= 0.6:
            m = _INLINE_DASH_RE.match(text)
            if m and m.group(1).strip() in inline_dash_labels:
                # "Anna—Why did you...?" -> em-dash cue + dialogue. Gated on a
                # recurring known speaker, so mid-sentence em-dashes and scene
                # headings ("Scene I.—...") don't trigger it.
                prev_type = _emit_cue(paras, para, m.group(1).strip(),
                                      m.group(2).strip())
                in_speech = True
                continue

        if ptype is None and inline_period_convention and _italic_ratio(para) <= 0.6:
            # "ANNA. Not a word more!"  or  "ANNA [coldly] Enough!" — ALL-CAPS named
            # speaker with a period or an immediate bracketed direction. Gated on
            # a recurring known speaker.
            label = remainder = None
            m = _INLINE_PERIOD_RE.match(text)
            if m and _is_allcaps_line(m.group(1).strip()) \
                    and m.group(1).strip() in inline_period_labels:
                label, remainder = m.group(1).strip(), m.group(2).strip()
            else:
                m2 = _INLINE_CAPS_DIR_RE.match(text)
                if m2 and m2.group(1).strip() in inline_period_labels:
                    label, remainder = m2.group(1).strip(), m2.group(2).strip()
            if label:
                prev_type = _emit_cue(paras, para, label, remainder)
                in_speech = True
                continue

        if ptype is None and _is_caps_cue(text):
            if prev_type == "Character":
                # A short all-caps line IMMEDIATELY after a cue is the speaker's
                # dialogue (an echo/exclamation like "AT LAST!"), not a second
                # cue or a scene heading.
                ptype = "Dialogue"
            elif cues_centered:
                centered = _is_centered(para)
                boldish = _bold_ratio(para) > 0.6
                if centered and not boldish:
                    ptype = "Character"
                else:
                    # Caps line breaking the cue convention: bold and/or
                    # left-aligned = scene title card / section header.
                    ptype = "Scene Heading"
                    if not boldish:
                        confident = False
            else:
                ptype = "Character"

        # Indent convention: classify by left-indent level when the author
        # uses distinct, consistent indents for dialogue vs. directions. This
        # is checked before the formatting fallbacks because indent is usually
        # the most reliable signal when present.
        if ptype is None and indent_convention:
            li = _round_indent(_left_indent(para))
            if abs(li - dialogue_indent) <= 0.08:
                ptype = "Dialogue"
            elif abs(li - direction_indent) <= 0.08:
                ptype = "Action"

        if ptype is None:
            is_wrapped = (text.startswith("(") and text.endswith(")")) or \
                         (text.startswith("[") and text.endswith("]"))
            if is_wrapped:
                # An interjected direction. Under a reliable parens convention a
                # cue opened this block and dialogue may follow, so it does NOT
                # end the speech: the next line is still that character's
                # dialogue. Speech is also preserved if the direction interrupts
                # immediately after a cue (the speaker has not spoken yet).
                ptype = "Action"
                runs = _docx_runs(para)
                if runs:
                    paras.append({"type": ptype, "runs": runs})
                    keep_speech = parens_convention or prev_type == "Character"
                    prev_type = ptype
                    if not keep_speech:
                        in_speech = False
                continue
            elif _italic_ratio(para) > 0.6 and not (in_speech and parens_convention):
                # Italic-dominant line = stage direction, UNLESS mid-speech in a
                # parens-convention doc (there italics mean emphasis and real
                # directions would be wrapped). In an italic-directions play
                # this is the primary, reliable direction signal.
                ptype = "Action"
                if italic_directions:
                    confident = True
                # If this direction interrupts immediately after a cue (the
                # speaker hasn't spoken yet), keep the speech open so the next
                # line is recognized as that character's dialogue.
                if prev_type == "Character":
                    runs = _docx_runs(para)
                    if runs:
                        paras.append({"type": ptype, "runs": runs})
                        prev_type = ptype
                        in_speech = True   # speech still pending after the aside
                    continue

        # Tier 3: context fallback
        if ptype is None:
            if prev_type == "Character":
                ptype = "Dialogue"            # confident: a cue precedes
            elif in_speech and (parens_convention or
                                prev_type == "Action"):
                ptype = "Dialogue"            # speech resumes after an interjected aside
            elif prev_type == "Dialogue":
                ptype = "Dialogue"            # unbroken speech block
            elif italic_directions and _italic_ratio(para) <= 0.6:
                # In an italic-directions play, a non-italic line that isn't a
                # cue is dialogue (directions would be italic).
                ptype = "Dialogue"
            else:
                ptype, confident = "Action", False   # truly unknown

        # Authors' manual line breaks (Shift+Enter) are junk in InDesign:
        # either split each visual line into its own paragraph of the same
        # style (default), or join them into one flowing paragraph.
        has_breaks = bool(re.search(r"[\n\x0b]", "".join(r.text for r in para.runs)))
        if has_breaks and breaks == "split":
            emitted = False
            for seg in _split_runs_on_breaks(para):
                seg = _finish(seg)
                if seg:
                    paras.append({"type": ptype, "runs": seg})
                    emitted = True
            if emitted and not confident:
                guessed.append((len(paras) - 1, ptype, text[:60]))
            if emitted:
                prev_type = ptype
                if ptype == "Character":
                    in_speech = True
                elif ptype == "Scene Heading":
                    in_speech = False
            continue

        runs = _docx_runs(para)
        if runs and breaks == "join":
            for r in runs:
                r["t"] = re.sub(r"[\n\x0b]+", " ", r["t"])
        if not runs:
            continue
        paras.append({"type": ptype, "runs": runs})
        if not confident:
            guessed.append((len(paras) - 1, ptype, text[:60]))
        prev_type = ptype
        if ptype == "Character":
            in_speech = True
        elif ptype == "Scene Heading":
            in_speech = False

    parse_word_docx.last_conventions = {
        "cues_centered": cues_centered,
        "parens_convention": parens_convention,
        "caps_cue_count": len(caps_centered),
        "caps_cue_centered_fraction": (
            round(sum(caps_centered) / len(caps_centered), 3)
            if caps_centered else None),
        "wrapped_dirs": wrapped_dirs,
        "unwrapped_italic_dirs": unwrapped_dirs,
        "italic_directions": italic_directions,
        "inline_cue_convention": inline_cue_convention,
        "inline_cue_labels": sorted(inline_cue_labels)[:8],
        "inline_dash_convention": inline_dash_convention,
        "inline_dash_labels": sorted(inline_dash_labels)[:8],
        "inline_period_convention": inline_period_convention,
        "inline_period_labels": sorted(inline_period_labels)[:8],
        "indent_convention": indent_convention,
        "dialogue_indent": dialogue_indent,
        "direction_indent": direction_indent,
        "is_musical": is_musical,
        "lyric_indent": lyric_indent,
        "cast_size": len(cast_names),
        "has_scene_headings": any(p["type"] == "Scene Heading" for p in paras),
    }
    return paras, guessed




def write_docx(paras, style_map, out_path, force_caps=True, keep_empty=False):
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        sys.exit("error: python-docx not installed — run:  pip install python-docx")

    doc = Document()
    # python-docx's default template has <w:zoom> without the required
    # w:percent attribute; patch it so output passes strict OOXML validation.
    from docx.oxml.ns import qn
    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")

    normal = doc.styles["Normal"]
    # Deliberately set NO font name. An explicit font imports into InDesign as
    # local character formatting and overrides the font defined in the matching
    # paragraph style (forcing e.g. Times New Roman). With no font set, the runs
    # carry none, so on Place the InDesign style definition supplies the font
    # automatically (no Clear Overrides step). A size is kept only so the .docx
    # is legible if opened in Word; size does not fight the imported style the
    # way a font name does.
    normal.font.size = Pt(12)

    # Cosmetic defaults so the docx itself is readable; InDesign replaces these
    # with the matching style definitions when the names match.
    COSMETICS = {
        "***Stage Directions":        dict(italic=True, indent=0.5),
        "***Character Name":          dict(center=True, space_before=10),
        "***Dialogue":                dict(),
        "***Song Lyrics":             dict(indent=0.5),
        "***Song Title":              dict(center=True, bold=True, space_before=12),
        "***Rap Lyrics":              dict(indent=0.5),
        "***Scene Number":            dict(bold=True, space_before=14),
        "***Scene Subtitle":          dict(italic=True, space_before=4),
        "FD General (front matter)":  dict(space_before=8),
    }

    created = {}

    def get_style(name):
        if name in created:
            return created[name]
        st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = normal
        st.quick_style = True
        cos = COSMETICS.get(name, {})
        if cos.get("italic"):
            st.font.italic = True
        if cos.get("bold"):
            st.font.bold = True
        if cos.get("center"):
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cos.get("indent"):
            st.paragraph_format.left_indent = Inches(cos["indent"])
        if cos.get("space_before"):
            st.paragraph_format.space_before = Pt(cos["space_before"])
        created[name] = st
        return st

    counts = Counter()
    for p in paras:
        ptype = p["type"]
        style_name = style_map.get(ptype, f"FD {ptype}")
        para = doc.add_paragraph(style=get_style(style_name))
        for r in p["runs"]:
            text = r["t"]
            if force_caps and ptype in UPPERCASE_TYPES:
                text = text.upper()
            run = para.add_run(re.sub(r"[\n\x0b]+", " ", text))
            run.italic = r["i"] or None
            run.bold = r["b"] or None
            run.underline = r["u"] or None
        counts[style_name] += 1

    # Strip every font reference from the styles part (docDefaults + each
    # style). python-docx's bundled template carries the Office theme
    # (Calibri/Cambria); if left in place, InDesign resolves text to those theme
    # fonts instead of the matching paragraph style's font. Removing all
    # <w:rFonts> leaves no font anywhere, so on Place the InDesign style
    # definition supplies the font cleanly.
    from docx.oxml.ns import qn
    styles_root = doc.styles.element
    for rfonts in styles_root.findall(".//" + qn("w:rFonts")):
        rfonts.getparent().remove(rfonts)

    doc.save(out_path)

    # Also neutralize the theme fonts inside the saved package, so nothing in
    # theme1.xml can reintroduce Cambria/Calibri as a fallback.
    _neutralize_theme_fonts(out_path)
    return counts


def _neutralize_theme_fonts(docx_path):
    """Rewrite theme1.xml so major/minor Latin typefaces are empty, removing
    the Office theme's Cambria/Calibri fallback that InDesign would otherwise
    resolve. Leaves the rest of the package untouched."""
    import zipfile, shutil, os, re as _re
    tmp = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path) as zin, \
         zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/theme/theme1.xml":
                text = data.decode("utf-8")
                # blank out the typeface on major/minor latin font slots
                text = _re.sub(r'(<a:latin[^>]*\btypeface=")[^"]*(")',
                               r"\1\2", text)
                data = text.encode("utf-8")
            zout.writestr(item, data)
    shutil.move(tmp, docx_path)


# ---------------------------------------------------------------------------
# Report: a short, plain-language summary of a run — what the manuscript looks
# like, what each paragraph was labeled, and which lines are worth a look. It
# quotes the real text so issues are easy to find in the source.
# ---------------------------------------------------------------------------

# Internal element type -> friendly label shown in the report.
_FRIENDLY = {
    "Character": "Character name",
    "Dialogue": "Dialogue",
    "Action": "Stage direction",
    "Scene Heading": "Scene number",
    "Lyrics": "Song lyric",
    "Song Title": "Song title",
    "Simultaneous": "Simultaneous dialogue",
    "Parenthetical": "Parenthetical",
    "General": "Front matter",
}


def _describe(conv):
    """One plain sentence describing the manuscript's apparent conventions."""
    kind = "musical" if conv.get("is_musical") else "play"
    if conv.get("inline_period_convention"):
        cue = "character names followed by a period"
    elif conv.get("inline_dash_convention"):
        cue = "character names joined to the dialogue with an em-dash"
    elif conv.get("inline_cue_convention"):
        cue = "character names followed by a colon"
    elif conv.get("cues_centered"):
        cue = "centered character names"
    else:
        cue = "character names on their own line"
    dirs = []
    if conv.get("parens_convention"):
        dirs.append("stage directions in parentheses")
    if conv.get("italic_directions"):
        dirs.append("stage directions in italics")
    if conv.get("indent_convention"):
        dirs.append("dialogue and directions set apart by indentation")
    sentence = f"Looks like a {kind} with {cue}"
    if dirs:
        sentence += ", and " + ", ".join(dirs)
    return sentence + "."


def build_report(path, breaks):
    """Return a short, plain-language report for a Word manuscript: what it
    looks like, what got labeled, and which lines are worth checking."""
    paras, guessed = parse_word_docx(path, breaks=breaks)
    conv = getattr(parse_word_docx, "last_conventions", {})

    def text_of(i):
        return "".join(r["t"] for r in paras[i]["runs"]).strip()

    def friendly(t):
        return _FRIENDLY.get(t, t)

    out = []
    out.append("=" * 68)
    out.append(f"autotypesetter report  ({os.path.basename(path)})")
    out.append("=" * 68)

    # What the manuscript looks like.
    out.append("\nWHAT IT LOOKS LIKE")
    out.append("  " + _describe(conv))

    # What got labeled (the part that worked).
    counts = Counter(p["type"] for p in paras if p["type"] != "Blank")
    total = sum(counts.values())
    out.append(f"\nWHAT IT LABELED  ({total} paragraphs)")
    for t, n in counts.most_common():
        out.append(f"  {n:>6}  {friendly(t)}")

    # Where it might be wrong: low-confidence guesses plus simple sanity checks,
    # collected per paragraph with the real text and a short reason.
    issues = {}
    for idx, _ptype, _snip in guessed:
        if 0 <= idx < len(paras):
            issues.setdefault(idx, "the tool was not sure about this line")
    for i, p in enumerate(paras):
        st = text_of(i)
        lbl = p["type"]
        if lbl == "Character" and len(st) > 35:
            issues[i] = "long for a character name — may be dialogue"
        elif lbl == "Dialogue" and st.startswith("(") and st.endswith(")"):
            issues[i] = "wrapped in parentheses — may be a stage direction"
        elif lbl == "Scene Heading" and not re.match(
                r"(?i)^\s*(scene|act|prologue|epilogue|french scene|tableau)\b", st):
            issues[i] = "labeled a scene number but doesn't read like one"
        elif (lbl == "Character" and i + 1 < len(paras)
              and paras[i + 1]["type"] not in ("Dialogue", "Action")):
            issues[i] = "a character name with no dialogue after it"

    out.append(f"\nLINES TO CHECK  ({len(issues)})")
    if not issues:
        out.append("  None — every paragraph was labeled with confidence.")
    else:
        for i in sorted(issues)[:40]:
            snippet = text_of(i)
            if len(snippet) > 75:
                snippet = snippet[:75] + "\u2026"
            out.append(f"  paragraph {i} \u2014 labeled \"{friendly(paras[i]['type'])}\"")
            out.append(f'      "{snippet}"')
            out.append(f"      ({issues[i]})")
        if len(issues) > 40:
            out.append(f"  \u2026and {len(issues) - 40} more")

    # Heads-up: a long opening run with no real cues is usually front matter.
    body_start = None
    for i in range(len(paras)):
        if (paras[i]["type"] == "Character" and i + 1 < len(paras)
                and paras[i + 1]["type"] == "Dialogue"):
            hits = sum(1 for j in range(i, min(i + 12, len(paras)))
                       if paras[j]["type"] == "Character"
                       and j + 1 < len(paras)
                       and paras[j + 1]["type"] == "Dialogue")
            if hits >= 3:
                body_start = i
                break
    if body_start and body_start > 5:
        # only worth mentioning if those paragraphs weren't already recognized
        # as front matter (a short preamble is auto-detected; a long preface is
        # not, and that is the case where --start-at-body helps)
        untrimmed = sum(1 for p in paras[:body_start] if p["type"] != "General")
        if untrimmed > 5:
            out.append("\nHEADS-UP")
            out.append(f"  The script seems to start around paragraph {body_start}; "
                       f"the {untrimmed} paragraphs before it look like a")
            out.append("  title page, cast list, or notes. Run again with "
                       "--start-at-body to drop them.")

    out.append("\n" + "=" * 68)
    out.append("If the labels look right and only the lines above need a look,")
    out.append("the file is ready to place in InDesign.")
    out.append("=" * 68)
    return "\n".join(out)


def _find_body_start(paras):
    """Find the paragraph index where the play body begins, dropping front
    matter (title, copyright, cast list, production notes). Handles plays that
    open on a scene marker or a run of stage directions, not just on dialogue.

    Strategy: find the first 'anchor' of real play content — a scene heading,
    or the first cue that begins a sustained cue/dialogue run — then walk
    backward past contiguous opening stage directions to include the scene's
    opening, but STOP at recognizable front-matter lines (cast list, notes).
    """
    def text_of(p):
        return "".join(r["t"] for r in p["runs"]).strip()

    # explicit scene heading wins — and IS the start; never back up past it.
    SCENE_WORD_RE = re.compile(
        r"(?i)^\s*(scene|act|prologue|epilogue|pre-?show|preshow|"
        r"curtain|tableau|french scene)\b")
    for i, p in enumerate(paras):
        if p["type"] == "Scene Heading" and SCENE_WORD_RE.match(text_of(p)):
            return i

    # in a musical with no scene headings, the first Song Title marks the body
    # (back up over an immediately preceding opening stage direction).
    first_song = next((i for i, p in enumerate(paras)
                       if p["type"] == "Song Title"), None)
    if first_song is not None:
        start = first_song
        k = first_song - 1
        while k >= 0 and paras[k]["type"] in ("Action",):
            tprev = text_of(paras[k])
            if re.search(r"(?i)(copyright|©|^by\s|note:|discrepan)", tprev):
                break
            start = k
            k -= 1
        return start

    # otherwise: first cue that starts a sustained cue->dialogue run
    def cue_dialogue_at(i):
        return (i + 1 < len(paras)
                and paras[i]["type"] == "Character"
                and paras[i + 1]["type"] == "Dialogue")
    anchor = None
    for i in range(len(paras)):
        if cue_dialogue_at(i) and sum(
                1 for j in range(i, min(i + 12, len(paras)))
                if cue_dialogue_at(j)) >= 3:
            anchor = i
            break
    if anchor is None:
        return 0

    # Front-matter lines the backup must never cross: cast-list entries
    # ("NAME: description" / "NAME\tdescription"), production notes, copyright,
    # title-ish lines. These often parse as Action/Dialogue, so they are
    # detected by text shape rather than type.
    FRONTMATTER_RE = re.compile(
        r"(?i)(^cast of characters|^dramatis personae|"
        r"\bnote:|\bnotes:|^casting\b|^setting\b|^time\b|^place\b|"
        r"copyright|^by\s|©|^rev\.|@)")

    def is_frontmatter(p):
        t = text_of(p)
        if not t:
            return False
        if FRONTMATTER_RE.search(t):
            return True
        # tabbed "NAME<tab>description" cast-list shape
        if "\t" in "".join(r["t"] for r in p["runs"]) and t[:1].isupper():
            return True
        return False

    # Walk backward from the anchor over contiguous stage directions (the
    # scene's opening staging), stopping at any front-matter line or a cue.
    start = anchor
    k = anchor - 1
    while k >= 0:
        p = paras[k]
        if p["type"] == "Character":
            break
        if is_frontmatter(p):
            break
        if p["type"] in ("Action", "Scene Heading", "Dialogue"):
            # only absorb directions/markers; a Dialogue here at the boundary
            # is usually a mis-typed opening direction (no cue context yet)
            start = k
            k -= 1
            continue
        break
    return start


def parse_song_list(src_path):
    """Extract the musical's song list from the source manuscript's front
    matter. Returns a list of (number, title, performers) tuples in order.
    Recognizes rows like '12.  Stuck   Andrew and Willow' or
    '2A. Pure Violation'. Stops once the list ends (body begins)."""
    from docx import Document
    doc = Document(src_path)
    rows = []
    row_re = re.compile(
        r"^\s*(\d+[A-Za-z]?)\.\s*\t?\s*([^\t]+?)(?:\t+(.+))?$")
    seen_numbers = set()
    for para in doc.paragraphs:
        t = para.text.rstrip()
        if not t.strip():
            continue
        m = row_re.match(t)
        if not m or len(t) > 90:
            # once collecting has started and a non-row appears, the list is done
            if rows:
                break
            continue
        num = m.group(1)
        if num in seen_numbers:        # looped into the body; stop
            break
        title = m.group(2).strip()
        who = (m.group(3) or "").strip().rstrip("\t").strip()
        # title-case-ish filter: skip obvious non-songs (very long "titles")
        if len(title) > 60:
            continue
        rows.append((num, title, who))
        seen_numbers.add(num)
    return rows


def _song_number_index(song_rows):
    """Map a normalized song title -> 'No. N' label, for body song titles."""
    idx = {}
    for num, title, _who in song_rows:
        key = re.sub(r"\s+", " ", title.strip().lower())
        idx[key] = f"No. {num}"
    return idx


def write_docx_from_template(paras, out_path, template_path, song_rows=None,
                             force_caps=True, type_to_style=None):
    """Write a fully-designed .docx by using `template_path` as a style donor:
    its styles, fonts, theme, borders and page setup are preserved, and the
    body is repopulated with our classified paragraphs mapped to the template's
    named styles. Author blank lines are preserved as empty Normal paragraphs
    so the source's vertical spacing is reproduced exactly."""
    from docx import Document
    from docx.oxml.ns import qn

    # element type -> template paragraph style name (overridable via config)
    TYPE_TO_STYLE = dict(DEFAULT_TEMPLATE_MAP)
    if type_to_style:
        TYPE_TO_STYLE.update(type_to_style)

    doc = Document(template_path)
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)

    available = {s.name for s in doc.styles}
    missing = {v for v in TYPE_TO_STYLE.values() if v not in available}
    if missing:
        sys.exit(f"error: template is missing required styles: {sorted(missing)}")
    blank_style = "Normal" if "Normal" in available else None

    song_idx = _song_number_index(song_rows) if song_rows else {}

    for p in paras:
        ptype = p["type"]

        if ptype == "Blank":
            doc.add_paragraph("", style=blank_style) if blank_style \
                else doc.add_paragraph("")
            continue

        text = "".join(r["t"] for r in p["runs"])
        style = TYPE_TO_STYLE.get(ptype, "Dialogue")

        if ptype == "Song Title":
            # reformat to "No. N | Title", pulling N from the source song list
            raw = text.strip()
            m = re.match(r"^\s*(\d+[A-Za-z]?)\.?\s*\t?\s*(.+)$", raw)
            title = m.group(2).strip() if m else raw
            key = re.sub(r"\s+", " ", title.strip().lower())
            label = song_idx.get(key)
            if label is None and m:
                label = f"No. {m.group(1)}"
            text = f"{label} | {title}" if label else title

        if ptype == "Character" and force_caps:
            text = text.upper()

        doc.add_paragraph(text, style=style)

    doc.save(out_path)
    return Counter(p["type"] for p in paras if p["type"] != "Blank")


# ---------------------------------------------------------------------------
# HTML input — a local .html/.htm file or a URL (e.g. a Project Gutenberg page)
#
# The HTML is parsed into block-level paragraphs, preserving italic/bold and
# any centering/indent, while dropping boilerplate (Project Gutenberg
# header/footer) and page-number spans. Those paragraphs are written to a .docx
# and run through the normal Word classifier, so HTML plays reuse every
# convention detector (centered cues, italic directions, em-dash/period inline
# cues, musical lyric indent, front-matter trimming) with no duplicate logic.
# ---------------------------------------------------------------------------

from html.parser import HTMLParser as _HTMLParser


class _PlayHTMLParser(_HTMLParser):
    BLOCK = {"p", "div", "h1", "h2", "h3", "h4", "h5", "h6",
             "li", "blockquote", "figcaption", "dd", "dt", "center"}
    DROP = {"head", "script", "style", "title", "nav"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.blocks = []
        self._cur = None
        self._stack = []          # (tag, kind) kind in i/b/drop/droptrack/block/None
        self._i = 0
        self._b = 0
        self._drop = 0

    @staticmethod
    def _attr(attrs, name):
        for k, v in attrs:
            if k == name:
                return (v or "")
        return ""

    @staticmethod
    def _indent(style):
        m = re.search(r"(?:margin-left|padding-left|text-indent)\s*:\s*"
                      r"([\d.]+)(em|px|pt|in)", style)
        if not m:
            return 0.0
        v, u = float(m.group(1)), m.group(2)
        return round({"em": v * 0.166, "px": v / 96.0,
                      "pt": v / 72.0, "in": v}.get(u, 0.0), 2)

    def handle_starttag(self, tag, attrs):
        cls = self._attr(attrs, "class").lower()
        style = self._attr(attrs, "style").lower()
        idv = self._attr(attrs, "id").lower()
        drop = (tag in self.DROP
                or any(x in cls for x in ("pg-boilerplate", "pgheader",
                                          "pgfooter", "pagen", "toc"))
                or idv.startswith("pg-header") or idv.startswith("pg-footer")
                or idv.startswith("page_") or idv.startswith("pageno")
                or "pagen" in idv)
        if drop:
            self._stack.append((tag, "drop"))
            self._drop += 1
            return
        if self._drop:
            self._stack.append((tag, "droptrack"))
            return
        if tag in ("i", "em", "cite"):
            self._stack.append((tag, "i"))
            self._i += 1
            return
        if tag in ("b", "strong"):
            self._stack.append((tag, "b"))
            self._b += 1
            return
        if tag == "br":
            if self._cur is not None:
                self._cur["runs"].append(("\n", False, False))
            return
        if tag == "hr":
            self._flush()
            return
        if tag in self.BLOCK:
            self._flush()
            center = ("center" in style or "center" in cls
                      or "centre" in cls or tag == "center")
            self._cur = {"runs": [], "center": center,
                         "indent": self._indent(style),
                         "heading": (len(tag) == 2 and tag[0] == "h")}
            self._stack.append((tag, "block"))
            return
        self._stack.append((tag, None))

    def handle_endtag(self, tag):
        while self._stack:
            t, kind = self._stack.pop()
            if kind == "i":
                self._i = max(0, self._i - 1)
            elif kind == "b":
                self._b = max(0, self._b - 1)
            elif kind == "drop":
                self._drop = max(0, self._drop - 1)
            elif kind == "block":
                self._flush()
            if t == tag:
                break

    def handle_data(self, data):
        if self._drop:
            return
        # Collapse source-formatting whitespace (including the newlines Project
        # Gutenberg wraps lines at) to single spaces, so a <p> stays one
        # paragraph. Real line breaks come only from <br>, inserted as "\n".
        data = re.sub(r"\s+", " ", data)
        if data == " " and (self._cur is None or not self._cur["runs"]):
            return
        if self._cur is None:
            if not data.strip():
                return
            self._cur = {"runs": [], "center": False,
                         "indent": 0.0, "heading": False}
        self._cur["runs"].append((data, self._i > 0, self._b > 0))

    def _flush(self):
        if self._cur and any(t.strip() for t, _, _ in self._cur["runs"]):
            self.blocks.append(self._cur)
        self._cur = None

    def close(self):
        super().close()
        self._flush()


def _read_html_source(src):
    """Return HTML text from a URL (http/https) or a local file path."""
    if re.match(r"^https?://", src, re.I):
        import urllib.request
        req = urllib.request.Request(
            src, headers={"User-Agent": "autotypesetter/1.0 (+manuscript typesetting)"})
        with urllib.request.urlopen(req, timeout=30) as r:
            raw = r.read()
        # honor a charset if present, else assume utf-8
        ctype = r.headers.get("Content-Type", "") if hasattr(r, "headers") else ""
        m = re.search(r"charset=([\w\-]+)", ctype, re.I)
        enc = m.group(1) if m else "utf-8"
        try:
            return raw.decode(enc, errors="replace")
        except LookupError:
            return raw.decode("utf-8", errors="replace")
    with open(src, encoding="utf-8", errors="replace") as fh:
        return fh.read()


def html_to_docx(src, out_path):
    """Convert an HTML play (URL or file) to a .docx, preserving italic/bold,
    centering and indentation, and dropping Project Gutenberg boilerplate and
    page-number markers. Returns out_path."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        sys.exit("error: python-docx not installed — run:  pip install python-docx")

    html = _read_html_source(src)

    # Trim Project Gutenberg boilerplate at the text level too, in case the
    # markup didn't carry the boilerplate classes.
    mstart = re.search(r"\*\*\*\s*START OF TH(E|IS) PROJECT GUTENBERG.*?\*\*\*",
                       html, re.I | re.S)
    if mstart:
        html = html[mstart.end():]
    mend = re.search(r"\*\*\*\s*END OF TH(E|IS) PROJECT GUTENBERG", html, re.I)
    if mend:
        html = html[:mend.start()]

    parser = _PlayHTMLParser()
    parser.feed(html)
    parser.close()

    doc = Document()
    from docx.oxml.ns import qn
    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")
    doc.styles["Normal"].font.size = Pt(12)

    for blk in parser.blocks:
        # merge adjacent runs of equal formatting; collapse internal whitespace
        merged = []
        for text, ital, bold in blk["runs"]:
            text = re.sub(r"[ \t\r\f\v]+", " ", text.replace("\xa0", " "))
            if not text:
                continue
            if merged and merged[-1][1] == ital and merged[-1][2] == bold:
                merged[-1][0] += text
            else:
                merged.append([text, ital, bold])
        line = "".join(t for t, _, _ in merged).strip()
        if not line:
            continue
        p = doc.add_paragraph()
        if blk["center"]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if blk["indent"] > 0.05:
            p.paragraph_format.left_indent = Inches(blk["indent"])
        for text, ital, bold in merged:
            run = p.add_run(text)
            if ital:
                run.italic = True
            if bold or blk["heading"]:
                run.bold = True

    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser(
        description="Turn a script manuscript (.fdx, .rtf, .docx, or .html/URL) "
                    "into an InDesign-ready .docx with named paragraph styles."
    )
    ap.add_argument("input", help="path to .fdx, .rtf, .docx, or .html (or an http/https URL)")
    ap.add_argument("-o", "--output", help="output .docx path (default: input name + .docx)")
    ap.add_argument(
        "--map", action="append", default=[], metavar="ELEMENT=STYLE",
        help='override/add a mapping, e.g. --map "Parenthetical=***Dialogue" (repeatable)',
    )
    ap.add_argument("--no-caps", action="store_true",
                    help="don't force-uppercase Character paragraphs")
    ap.add_argument("--keep-empty", action="store_true",
                    help="keep empty paragraphs instead of dropping them")
    ap.add_argument("--breaks", choices=["split", "join"], default="split",
                    help="manual line breaks in Word input: 'split' each "
                         "visual line into its own paragraph (default) or "
                         "'join' them into one paragraph with spaces")
    ap.add_argument("--report", action="store_true",
                    help="don't convert; print a short, plain-language report of "
                         "what the manuscript looks like, what each paragraph was "
                         "labeled, and which lines are worth checking")
    ap.add_argument("--start-at-body", action="store_true",
                    help="drop all front matter (title page, notes, character "
                         "list) and begin the output at the first scene heading "
                         "or, if there is none, the first character cue")
    ap.add_argument("--raw-docx", metavar="RAW.docx",
                    help="for HTML input: path for the plain, unclassified "
                         "author-style .docx (the pre-parse copy). Defaults to "
                         "<name>_raw.docx next to the output.")
    ap.add_argument("--raw-only", action="store_true",
                    help="for HTML input: only produce the raw author-style "
                         ".docx (a copy-paste-style manuscript for building a "
                         "test corpus) and skip classification entirely.")
    ap.add_argument("--no-raw", action="store_true",
                    help="for HTML input: do NOT save the raw author-style "
                         ".docx (classify via a throwaway temp file instead).")
    ap.add_argument("--config", metavar="CONFIG.json",
                    help="JSON file of element->style mappings, overriding the "
                         "built-in defaults. May contain an \"indesign\" and/or "
                         "\"template\" object (a flat object is treated as "
                         "InDesign overrides). Only the keys present are changed.")
    ap.add_argument("--template", metavar="TEMPLATE.docx",
                    help="produce a fully-designed Word document using the given "
                         ".docx as a style donor (its fonts, named styles, song-"
                         "title box, page setup are reproduced). Maps elements to "
                         "the template's CHARACTER NAME / Dialogue / LYRICS / "
                         "Stage Direction / Song Titles / SCENE styles, and "
                         "preserves the author's blank-line spacing.")
    args = ap.parse_args()

    # Base maps come from defaults, optionally overridden by a --config file,
    # then by any inline --map flags (which take final precedence).
    if args.config:
        style_map, template_map = load_config(args.config)
    else:
        style_map = dict(DEFAULT_STYLE_MAP)
        template_map = dict(DEFAULT_TEMPLATE_MAP)
    for spec in args.map:
        if "=" not in spec:
            sys.exit(f'error: bad --map "{spec}" (expected ELEMENT=STYLE)')
        k, v = spec.split("=", 1)
        style_map[k.strip()] = v.strip()

    path = args.input
    lower = path.lower()
    guessed = []

    # HTML input (local .html/.htm file or an http/https URL). The HTML is first
    # converted to a plain, unclassified "author-style" .docx: just the text
    # with bold/italic/centering/indent preserved and no paragraph styles, as if
    # an author had pasted the page into Word. That raw .docx is what gets
    # parsed, and unless suppressed it is also saved as a reusable test
    # manuscript. Classification runs on it afterwards.
    if (lower.endswith(".html") or lower.endswith(".htm")
            or re.match(r"^https?://", path, re.I)):
        base = re.sub(r"[?#].*$", "", path).rstrip("/").split("/")[-1]
        base = re.sub(r"\.html?$", "", base, flags=re.I) or "play"
        # Where to write the raw author-style .docx (the pre-parse copy).
        raw_path = args.raw_docx or (base + "_raw.docx")
        if not args.no_raw or args.raw_only:
            html_to_docx(path, raw_path)
            print(f"wrote raw author-style manuscript: {raw_path}")
            parse_path = raw_path
        else:
            import tempfile
            parse_path = os.path.join(tempfile.gettempdir(),
                                      "autotypesetter_html.docx")
            html_to_docx(path, parse_path)
        if args.raw_only:
            # conversion-only mode: build the training manuscript and stop
            return
        if not (args.output or args.report):
            args.output = base + "_styled.docx"
        path = parse_path
        lower = ".docx"

    if args.report:
        if not lower.endswith(".docx"):
            sys.exit("error: --report is for Word (.docx) manuscripts "
                     "(.fdx/.rtf classify deterministically and need no report)")
        print(build_report(path, breaks=args.breaks))
        return

    if lower.endswith(".fdx"):
        paras, dual_blocks = parse_fdx(path)
    elif lower.endswith(".rtf"):
        paras, dual_blocks = parse_fd_rtf(path)
    elif lower.endswith(".docx"):
        paras, guessed = parse_word_docx(path, breaks=args.breaks,
                                         keep_blanks=bool(args.template))
        dual_blocks = []
    else:
        sys.exit("error: input must be .fdx, .rtf, .docx, .html/.htm, or an http(s) URL")

    trimmed_count = 0
    if args.start_at_body:
        start = _find_body_start(paras)
        if start:
            trimmed_count = start
            paras = paras[start:]
            guessed = [(idx - start, t, s) for idx, t, s in guessed
                       if idx >= start]
            dual_blocks = [(idx - start, w) for idx, w in dual_blocks
                           if idx >= start]
    # drop any leading blank markers so the body starts cleanly
    while paras and paras[0]["type"] == "Blank":
        paras = paras[1:]

    out = args.output or re.sub(r"\.(fdx|rtf|docx)$", "", path, flags=re.I) + "_styled.docx"

    if args.template:
        song_rows = []
        if lower.endswith(".docx"):
            try:
                song_rows = parse_song_list(path)
            except Exception:
                song_rows = []
        counts = write_docx_from_template(
            paras, out, args.template,
            song_rows=song_rows, force_caps=not args.no_caps,
            type_to_style=template_map)
        print(f"wrote {out}  (designed from template {os.path.basename(args.template)})")
        if trimmed_count:
            print(f"trimmed {trimmed_count} front-matter paragraph(s)")
        if song_rows:
            print(f"song list: {len(song_rows)} numbers parsed")
        print(f"{sum(counts.values())} paragraphs:")
        for name, n in counts.most_common():
            print(f"  {n:>6}  {name}")
        if guessed:
            print(f"\n{len(guessed)} paragraph(s) classified by context only — "
                  f"review these:")
            for idx, ptype, snippet in guessed[:25]:
                print(f"  #{idx:<5} guessed {ptype:<9} | {snippet}")
            if len(guessed) > 25:
                print(f"  ... and {len(guessed) - 25} more")
        return

    counts = write_docx(paras, style_map, out,
                        force_caps=not args.no_caps,
                        keep_empty=args.keep_empty)

    print(f"wrote {out}")
    if trimmed_count:
        print(f"trimmed {trimmed_count} front-matter paragraph(s); "
              f"output begins at the script body")
    print(f"{sum(counts.values())} paragraphs:")
    for name, n in counts.most_common():
        flag = "   <-- unmapped, restyle in InDesign" if name.startswith("FD ") else ""
        print(f"  {n:>6}  {name}{flag}")
    if dual_blocks:
        print(f"\n{len(dual_blocks)} dual-dialogue block(s) flattened to sequential "
              f"speech (hand-typeset if you want them side-by-side):")
        for idx, who in dual_blocks:
            print(f"  near paragraph {idx} (starts with {who})")
    if guessed:
        print(f"\n{len(guessed)} paragraph(s) classified by context only — "
              f"review these in InDesign:")
        for idx, ptype, snippet in guessed[:25]:
            print(f"  #{idx:<5} guessed {ptype:<9} | {snippet}")
        if len(guessed) > 25:
            print(f"  ... and {len(guessed) - 25} more")


if __name__ == "__main__":
    main()
